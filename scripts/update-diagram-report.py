from __future__ import annotations

import argparse
import datetime as dt
import os
import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


APP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def remove_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def copy_row_format(source_row, target_row) -> None:
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        target_tc_pr = target_cell._tc.tcPr
        target_cell._tc.remove(target_tc_pr)
        target_cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))
        if source_cell.paragraphs and target_cell.paragraphs:
            target_p = target_cell.paragraphs[0]._p
            target_p_pr = target_p.get_or_add_pPr()
            target_p.remove(target_p_pr)
            target_p.insert(0, deepcopy(source_cell.paragraphs[0]._p.get_or_add_pPr()))


def patch_app_properties(path: Path, pages: int, words: int) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as handle:
        temp_path = Path(handle.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w") as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "docProps/app.xml":
                    root = ET.fromstring(data)
                    pages_node = root.find(f"{{{APP_NS}}}Pages")
                    words_node = root.find(f"{{{APP_NS}}}Words")
                    if pages_node is not None:
                        pages_node.text = str(pages)
                    if words_node is not None:
                        words_node.text = str(words)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                target.writestr(item, data)
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def count_words(document: Document) -> int:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return len(re.findall(r"\b[\wÀ-ỹ]+\b", "\n".join(chunks), flags=re.UNICODE))


def update_document(input_path: Path, output_path: Path, pages: int) -> None:
    document = Document(input_path)

    paragraph_updates = {
        4: (
            "Repository: D:\\Users\\User\\IdeaProjects\\VibeGraph\n"
            "Evidence snapshot: 2026-08-14T10:12:42+07:00\n"
            "HEAD: d5154c4c368d7ca89fabb8da91a79858bea7af7b\n"
            "Generated from the current shared working-tree baseline"
        ),
        18: (
            "Bộ sơ đồ cập nhật được xây dựng từ source/configuration, migration, GitNexus và "
            "database runtime hiện tại. VibeGraph dùng Spring Boot/Vue, PostgreSQL cho control "
            "plane, Neo4j cho code graph, phân tích Java đồng bộ hoặc bất đồng bộ tùy route, "
            "watcher incremental qua STOMP, và MCP streamable HTTP có API-key/project binding."
        ),
        20: (
            "Quá trình cập nhật tài liệu không sửa production code. Worktree vẫn có thay đổi song "
            "song từ các session khác; watcher ghi nhận drift và không tự suy diễn hoặc tự sửa "
            "canonical diagram."
        ),
        34: (
            "Container quan sát: vibegraph-postgres, postgres:16.11-alpine, host "
            "127.0.0.1:5433, trạng thái healthy. Flyway có 19 migration thành công; schema hiện "
            "có 21 domain tables, 23 foreign keys và 66 domain-table indexes (68 public indexes "
            "khi tính hai index của Flyway)."
        ),
        36: (
            "Các UUID actor/target của audit_logs là logical references, không phải foreign keys, "
            "sau V15__audit_log_transaction_hardening.sql. Volatile table row counts được loại khỏi "
            "canonical ERD; schema counts mới là bằng chứng ổn định hơn."
        ),
        38: (
            "Container quan sát: vibegraph-neo4j, neo4j:5.26-community, Bolt "
            "127.0.0.1:7687, Browser 127.0.0.1:7474, database neo4j online. V1/V2 Cypher "
            "migrations tạo label-scoped constraints/indexes và shared Symbol indexes."
        ),
        42: (
            "docker-compose.yml chứng minh frontend nginx/Vue, backend Spring Boot/Java 21, "
            "PostgreSQL, Neo4j, health dependencies và writable /app/projects + /app/uploads "
            "mounts. Parser/analysis, graph và các MCP tool cần graph đi qua GraphRepository; chỉ "
            "Neo4jGraphRepository giữ raw Neo4j Driver."
        ),
        44: (
            "Authentication: register tạo User/settings mới; login xác thực User hiện hữu; OAuth "
            "callback link/tạo identity. AuthController hoặc OAuth2LoginSuccessHandler dùng "
            "AuthCookieService để set/replace/clear cookies; RefreshSessionService issue/rotate/revoke."
        ),
        45: (
            "Import/analyze: archive mặc định phân tích đồng bộ và trả 200; archive async=true, "
            "GitHub và local submit analysisExecutor trước khi controller trả 202. Manual analyze "
            "là luồng riêng qua ProjectAnalysisScheduler."
        ),
        46: (
            "Watcher: getFileSlice(before) -> deleteFile -> parse/upsert nếu file còn tồn tại -> "
            "getFileSlice(after) -> tính delta -> broadcast incremental. Source không chứng minh "
            "rollback transaction cho toàn bộ chuỗi."
        ),
        47: (
            "Graph/source/impact thuộc browser flow. Local patch là endpoint CLI/JWT hoặc "
            "project-bound API key; sau commit, PatchAnalysisScheduler coalesce và schedule full "
            "background re-analysis."
        ),
        48: (
            "Use-case UML: response chứa actors/use cases với source/confidence, relations, inference "
            "warnings, PlantUML/Mermaid và projected views; frontend render sanitized SVG và tải PNG."
        ),
        58: (
            "GitNexus index được xác nhận up-to-date tại HEAD: 1,173 files, 17,907 symbols, "
            "41,198 relationships và 300 execution flows."
        ),
        61: "Frontend unit suite pass: 67 files, 570 tests.",
        63: (
            "PlantUML marker/canonical-copy validation và diagrams.net XML validation pass; page "
            "inventory là 10/6/2/1/2, không duplicate ID, broken edge reference hoặc out-of-page vertex."
        ),
    }
    for index, text in paragraph_updates.items():
        set_paragraph_text(document.paragraphs[index], text)

    for index in range(7, 15):
        paragraph = document.paragraphs[index]
        set_paragraph_text(paragraph, re.sub(r"^\d+\.\s*", "", paragraph.text))

    baseline = document.tables[3]
    set_cell_text(baseline.cell(2, 1), "2026-08-14T10:12:42+07:00 (Asia/Bangkok)")
    set_cell_text(baseline.cell(4, 1), "1,173 files; 17,907 symbols; 41,198 relationships; 300 execution flows")
    set_cell_text(baseline.cell(5, 1), "Targeted backend audit suite: 78 tests PASS; compile PASS")
    set_cell_text(baseline.cell(6, 1), "npm run type-check: PASS; unit: 67 files / 570 tests PASS")
    set_cell_text(baseline.cell(7, 1), "Current shared worktree; later drift is recorded by the watcher and requires review")

    capabilities = document.tables[4]
    set_cell_text(capabilities.cell(1, 1), "Register creates a local account; login authenticates an existing account; OAuth links/creates a verified identity; rotating refresh sessions and controller-owned cookies.")
    set_cell_text(capabilities.cell(2, 1), "Create/list/get/trash/restore/purge; archive 200 by default, archive async=true/GitHub/local 202; CLI setup and local patch endpoints.")
    set_cell_text(capabilities.cell(3, 1), "Import services submit their executor before 202; manual analyze is separately coalesced by ProjectAnalysisScheduler.")
    set_cell_text(capabilities.cell(4, 1), "Exact per-file order: before slice, delete, optional parse/upserts, after slice, delta, incremental broadcast.")
    set_cell_text(capabilities.cell(5, 1), "Browser graph/source/impact; CLI or project-bound API-key patch with scheduled full re-analysis.")
    set_cell_text(capabilities.cell(6, 1), "Only use-case endpoint; response includes source/confidence, inference warnings, relations, PlantUML/Mermaid and projected views.")
    set_cell_text(capabilities.cell(8, 1), "Users/plans/flags, credits/pricing, audit/security streams, request-abuse views, IP-block CRUD, support and read-only storage overview.")

    pg_table = document.tables[5]
    set_cell_text(pg_table.cell(3, 1), "66 domain-table indexes; 68 public-schema indexes including two Flyway metadata indexes")
    set_cell_text(pg_table.cell(4, 0), "Volatile row counts")
    set_cell_text(pg_table.cell(4, 1), "Intentionally omitted from canonical ERD; re-query and timestamp if operational counts are needed")

    neo_table = document.tables[6]
    set_cell_text(neo_table.cell(3, 1), "CONTAINS 2,786; DEFINES 5,010; HAS_METHOD 12,043; HAS_FIELD 9,483; CALLS 8,801; IMPORTS 8,742; READS 36,861; WRITES 14,469; INSTANTIATES 3,154; ANNOTATED_BY 1,712 legacy; HANDLES_ROUTE 620; STEP_IN_FLOW 284")
    set_cell_text(neo_table.cell(4, 0), "Schema/runtime distinctions")
    set_cell_text(neo_table.cell(4, 1), "Route has V1 constraint/index but runtime count 0; current parser emits APIEndpoint (620). isStub and Neo4j Project.status are not current runtime properties.")

    comparison = document.tables[8]
    set_cell_text(comparison.cell(3, 1), "Archive upload, GitHub tarball and local import only; archive is synchronous by default.")
    set_cell_text(comparison.cell(8, 1), "Current endpoints include APIEndpoint and corrected IMPORTS/HAS_INNER/OVERRIDES/HAS_RELATION/INJECTS directions; ANNOTATED_BY is documented as persisted legacy data.")

    artifacts = document.tables[9]
    set_cell_text(artifacts.cell(1, 1), "Evidence policy, canonical map, generators/watcher and deliberate exclusions.")
    set_cell_text(artifacts.cell(2, 1), "Current evidence identity, build/test results and timestamped PostgreSQL/Neo4j facts; live/ keeps rolling state separately.")
    set_cell_text(artifacts.cell(6, 0), "Combined PlantUML mirror")
    set_cell_text(artifacts.cell(6, 1), "VibeGraph_All_PlantUML_Diagrams.md, generated exactly from the three canonical sources.")
    set_cell_text(artifacts.cell(8, 1), "Four family-specific old-to-current records plus CHANGES-FROM-OLD.md cross-family index.")
    script_row = next((row for row in artifacts.rows if row.cells[0].text == "scripts/*diagram*.ps1"), None)
    if script_row is None:
        script_row = artifacts.add_row()
        copy_row_format(artifacts.rows[-2], script_row)
    set_cell_text(script_row.cells[0], "scripts/*diagram*.ps1")
    set_cell_text(script_row.cells[1], "Regenerate diagrams.net, synchronize combined PlantUML, and continuously capture evidence drift without editing canonical artifacts.")

    set_cell_text(
        document.tables[10].cell(0, 0),
        "XML validation confirms 10 use-case pages, 6 activity pages, 2 ERD pages, 1 component/deployment page and 2 class pages. The new files keep the reviewed page inventory while replacing stale content with evidence-backed models.",
    )

    sources = document.tables[12]
    set_cell_text(sources.cell(3, 1), "Register/login/refresh/logout/me and controller-owned cookie handling.")
    set_cell_text(sources.cell(8, 1), "Exact watcher before/delete/parse-upsert/after/broadcast order and non-atomic limitation.")
    set_cell_text(sources.cell(9, 1), "Only evidenced UML use-case endpoint and real response contract.")
    set_cell_text(sources.cell(13, 1), "Deployment, health dependencies, writable mounts, STOMP and optional realtime/high-volume storage.")

    if len(document.paragraphs) > 73:
        delete_paragraph(document.paragraphs[73])

    properties = document.core_properties
    properties.title = "VibeGraph - Báo cáo kiến trúc và đối chiếu bằng chứng"
    properties.subject = "Verified diagram update aligned with current code, migrations and local database runtime"
    properties.author = "VibeGraph documentation audit"
    properties.last_modified_by = "VibeGraph documentation audit"
    properties.comments = (
        "Evidence-backed VibeGraph diagram update aligned with current code, migrations "
        "and local runtime."
    )
    properties.created = dt.datetime(2026, 8, 14, 10, 12, 42)
    properties.modified = dt.datetime.now().replace(microsecond=0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if input_path.resolve() == output_path.resolve():
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=output_path.parent) as handle:
            staged = Path(handle.name)
        document.save(staged)
        os.replace(staged, output_path)
    else:
        document.save(output_path)

    saved = Document(output_path)
    patch_app_properties(output_path, pages=pages, words=count_words(saved))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--pages", type=int, default=7)
    args = parser.parse_args()
    update_document(args.input, args.output or args.input, args.pages)


if __name__ == "__main__":
    main()
